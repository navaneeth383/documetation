import streamlit as st
import docx
from docx.shared import Inches
import os
import tempfile
from io import BytesIO
import nbformat
from nbconvert import PythonExporter
import re
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

st.set_page_config(page_title="📘 Project Documentation Generator", layout="wide")

st.title("📘 Project Documentation Generator")
st.markdown("""
### Created by Gattu Navaneeth Rao
Upload your project code, screenshots, and supporting files to generate a detailed project-level documentation.
""")

st.markdown("---")

uploaded_files = st.file_uploader(
    "📂 Upload Project Files (.py, .sql, .ipynb, .txt, images, excels, etc)",
    accept_multiple_files=True
)

custom_notes = st.text_area("📝 Additional Notes to Include in Document")

generate_button = st.button("📄 Generate Documentation")

def detect_code_type(code):
    if "SELECT" in code.upper() or "CREATE TABLE" in code.upper():
        return "SQL"
    elif "def " in code or "import " in code:
        return "Python"
    else:
        return "Text"

def clean_code(text):
    return re.sub(r'\n+', '\n', text.strip())

def extract_from_ipynb(uploaded):
    content = uploaded.read()
    nb = nbformat.reads(content.decode("utf-8"), as_version=4)
    exporter = PythonExporter()
    (body, _) = exporter.from_notebook_node(nb)
    return body

def safe_add_code(doc, content):
    cleaned = clean_code(content)
    for block in cleaned.split('\n'):
        try:
            doc.add_paragraph(block)
        except Exception:
            doc.add_paragraph("[Error displaying this line due to encoding issues]")

if generate_button and uploaded_files:
    doc = docx.Document()
    doc.add_heading("Project Documentation", 0)
    doc.add_paragraph("Author: Gattu Navaneeth Rao")
    doc.add_paragraph("\n")
    doc.add_paragraph("Table of Contents")
    doc.add_paragraph("\n")

    for uploaded in uploaded_files:
        filename = uploaded.name
        extension = os.path.splitext(filename)[-1].lower()

        if extension in ['.png', '.jpg', '.jpeg']:
            with tempfile.NamedTemporaryFile(delete=False, suffix=extension) as tmp_file:
                tmp_file.write(uploaded.read())
                tmp_path = tmp_file.name
            doc.add_heading(f"Screenshot: {filename}", level=1)
            doc.add_picture(tmp_path, width=Inches(5.5))
            os.unlink(tmp_path)
        elif extension == '.ipynb':
            doc.add_heading(f"Notebook: {filename}", level=1)
            content = extract_from_ipynb(uploaded)
            safe_add_code(doc, content)
        else:
            content = uploaded.read().decode(errors='ignore')
            code_type = detect_code_type(content)
            doc.add_heading(f"{code_type} File: {filename}", level=1)
            safe_add_code(doc, content)

    if custom_notes:
        doc.add_page_break()
        doc.add_heading("Additional Notes", level=1)
        doc.add_paragraph(custom_notes)

    doc.add_page_break()
    doc.add_paragraph("\nGenerated using Streamlit by Gattu Navaneeth Rao")

    output = BytesIO()
    doc.save(output)
    output.seek(0)

    st.success("📄 Documentation generated successfully!")
    st.download_button(
        label="📥 Download Project_Documentation.docx",
        data=output,
        file_name="Project_Documentation.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
else:
    st.info("Upload at least one file and click Generate Documentation to start.")
